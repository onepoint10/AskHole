"""
File Manager Module
Handles file operations, uploads, downloads, and media processing
"""

import os
import shutil
import tempfile
import mimetypes
from pathlib import Path
from typing import List, Optional, Tuple, Dict, Any
from datetime import datetime
import json
import wave
import threading
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
import pygame
import io
import tkinterdnd2 as tkdnd
import subprocess
import sys
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.utils import ImageReader
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as ReportLabImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import markdown
import docx


class FileManager:
    """Manages file operations for the application"""
    
    def __init__(self, config_manager):
        """Initialize file manager"""
        self.config_manager = config_manager
        self.temp_files = []  # Track temporary files for cleanup
        self.downloads_dir = self._get_downloads_directory()
        
        # Initialize pygame mixer for audio playback
        try:
            pygame.mixer.init()
            self.audio_enabled = True
        except:
            self.audio_enabled = False
            print("Warning: Audio playback not available")
    
    def _get_downloads_directory(self) -> Path:
        """Get default downloads directory"""
        downloads_dir = Path.home() / "Downloads" / "GeminiClient"
        downloads_dir.mkdir(parents=True, exist_ok=True)
        return downloads_dir
    
    def select_files(self, multiple=True, file_types=None) -> List[str]:
        """Open file dialog to select files"""
        if file_types is None:
            file_types = [
                ("All supported", "*.pdf;*.txt;*.docx;*.xlsx;*.pptx;*.jpg;*.jpeg;*.png;*.gif;*.webp;*.mp3;*.wav;*.ogg"),
                ("Documents", "*.pdf;*.txt;*.docx;*.xlsx;*.pptx;*.csv;*.rtf"),
                ("Images", "*.jpg;*.jpeg;*.png;*.gif;*.webp;*.heic;*.heif"),
                ("Audio", "*.mp3;*.wav;*.ogg;*.m4a"),
                ("All files", "*.*")
            ]
        
        initial_dir = self.config_manager.get("last_used_directory", str(Path.home()))
        
        if multiple:
            files = filedialog.askopenfilenames(
                title="Select files",
                initialdir=initial_dir,
                filetypes=file_types
            )
        else:
            file = filedialog.askopenfilename(
                title="Select file",
                initialdir=initial_dir,
                filetypes=file_types
            )
            files = [file] if file else []
        
        if files:
            # Update last used directory
            self.config_manager.set("last_used_directory", os.path.dirname(files[0]))
        
        return list(files)
    
    def select_save_location(self, default_name="", file_types=None) -> str:
        """Open save dialog to select save location"""
        if file_types is None:
            file_types = [("Text files", "*.txt"), ("All files", "*.*")]
        
        initial_dir = self.config_manager.get("last_used_directory", str(Path.home()))
        
        file_path = filedialog.asksaveasfilename(
            title="Save as",
            initialdir=initial_dir,
            initialfile=default_name,
            filetypes=file_types
        )
        
        if file_path:
            self.config_manager.set("last_used_directory", os.path.dirname(file_path))
        
        return file_path
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get file information"""
        if not os.path.exists(file_path):
            return {}
        
        stat = os.stat(file_path)
        mime_type, _ = mimetypes.guess_type(file_path)
        
        return {
            "path": file_path,
            "name": os.path.basename(file_path),
            "size": stat.st_size,
            "size_str": self._format_file_size(stat.st_size),
            "mime_type": mime_type or "unknown",
            "modified": datetime.fromtimestamp(stat.st_mtime),
            "is_image": self._is_image_file(file_path),
            "is_audio": self._is_audio_file(file_path),
            "is_document": self._is_document_file(file_path)
        }
    
    def _format_file_size(self, size: int) -> str:
        """Format file size in human readable format"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024.0:
                return f"{size:.1f} {unit}"
            size /= 1024.0
        return f"{size:.1f} TB"
    
    def _is_image_file(self, file_path: str) -> bool:
        """Check if file is an image"""
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type and mime_type.startswith('image/')
    
    def _is_audio_file(self, file_path: str) -> bool:
        """Check if file is an audio file"""
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type and mime_type.startswith('audio/')
    
    def _is_document_file(self, file_path: str) -> bool:
        """Check if file is a document"""
        mime_type, _ = mimetypes.guess_type(file_path)
        if not mime_type:
            return False
        
        document_types = [
            'application/pdf',
            'text/plain',
            'application/vnd.openxmlformats-officedocument',
            'application/msword',
            'application/vnd.ms-excel',
            'application/vnd.ms-powerpoint',
            'text/csv',
            'application/rtf'
        ]
        
        return any(mime_type.startswith(doc_type) for doc_type in document_types)

    def _is_pdf_file(self, file_path: str) -> bool:
        """Check if file is a PDF"""
        return file_path.lower().endswith('.pdf')
    
    def create_thumbnail(self, image_path: str, size=(150, 150)) -> Optional[ImageTk.PhotoImage]:
        """Create thumbnail for image file"""
        try:
            with Image.open(image_path) as img:
                # Convert to RGB if necessary
                if img.mode in ('RGBA', 'LA', 'P'):
                    img = img.convert('RGB')
                
                # Create thumbnail
                img.thumbnail(size, Image.Resampling.LANCZOS)
                
                # Convert to PhotoImage
                return ImageTk.PhotoImage(img)
        except Exception as e:
            print(f"Error creating thumbnail: {e}")
            return None
    
    def save_image(self, image: Image.Image, filename: str = None) -> str:
        """Save PIL Image to file"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"generated_image_{timestamp}.png"
        
        file_path = self.downloads_dir / filename
        
        try:
            image.save(file_path, 'PNG')
            return str(file_path)
        except Exception as e:
            raise Exception(f"Failed to save image: {e}")
    
    def save_audio(self, audio_data: io.BytesIO, filename: str = None) -> str:
        """Save audio data to file"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"generated_audio_{timestamp}.wav"
        
        file_path = self.downloads_dir / filename
        
        try:
            with open(file_path, 'wb') as f:
                f.write(audio_data.getvalue())
            return str(file_path)
        except Exception as e:
            raise Exception(f"Failed to save audio: {e}")
    
    def save_text(self, text: str, filename: str = None) -> str:
        """Save text to file"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"response_{timestamp}.txt"
        
        file_path = self.downloads_dir / filename
        
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return str(file_path)
        except Exception as e:
            raise Exception(f"Failed to save text: {e}")
    
    def play_audio(self, audio_path: str, callback=None):
        """Play audio file"""
        if not self.audio_enabled:
            if callback:
                callback("Audio playback not available")
            return
        
        def play_thread():
            try:
                pygame.mixer.music.load(audio_path)
                pygame.mixer.music.play()
                
                # Wait for playback to finish
                while pygame.mixer.music.get_busy():
                    threading.Event().wait(0.1)
                
                if callback:
                    callback(None)
            except Exception as e:
                if callback:
                    callback(str(e))
        
        thread = threading.Thread(target=play_thread)
        thread.daemon = True
        thread.start()
    
    def stop_audio(self):
        """Stop audio playback"""
        if self.audio_enabled:
            pygame.mixer.music.stop()
    
    def create_temp_file(self, suffix="", prefix="gemini_", content=None) -> str:
        """Create temporary file"""
        temp_file = tempfile.NamedTemporaryFile(
            suffix=suffix, 
            prefix=prefix, 
            delete=False
        )
        
        if content:
            if isinstance(content, str):
                temp_file.write(content.encode('utf-8'))
            else:
                temp_file.write(content)
        
        temp_file.close()
        self.temp_files.append(temp_file.name)
        return temp_file.name
    
    def cleanup_temp_files(self):
        """Clean up temporary files"""
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except:
                pass
        self.temp_files.clear()

    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """Validate file for upload"""
        if not os.path.exists(file_path):
            return False, "File does not exist"

        file_size = os.path.getsize(file_path)
        if file_size > 20 * 1024 * 1024:  # 20MB limit
            return False, "File size exceeds 20MB limit"

        if file_size == 0:
            return False, "File is empty"

        # Check if file type is supported
        mime_type, _ = mimetypes.guess_type(file_path)
        file_extension = os.path.splitext(file_path)[1].lower()

        supported_types = [
            # PDF support for OpenRouter
            'application/pdf',
            # Text files
            'text/plain',
            'text/csv',
            'text/html',
            'text/css',
            'text/javascript',
            'application/x-javascript',
            'text/x-typescript',
            'application/json',
            'text/xml',
            'application/rtf',
            'text/rtf',
            # Office documents
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/vnd.openxmlformats-officedocument.presentationml.presentation',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/msword',
            'application/vnd.ms-excel',
            'application/vnd.ms-powerpoint',
            # Images
            'image/jpeg',
            'image/png',
            'image/gif',
            'image/webp',
            'image/heic',
            'image/heif',
            # Audio
            'audio/mpeg',
            'audio/wav',
            'audio/ogg'
        ]

        # Additional check for file extensions that might not have proper MIME types
        supported_extensions = ['.md', '.py', '.txt', '.csv', '.html', '.css', '.js', '.json', '.xml', '.pdf']

        if mime_type in supported_types or file_extension in supported_extensions:
            return True, "File is valid"
        else:
            return False, f"Unsupported file type: {mime_type or file_extension}"

    def can_convert_to_pdf(self, file_path: str) -> bool:
        """Check if file can be converted to PDF"""
        file_extension = os.path.splitext(file_path)[1].lower()
        convertible_extensions = [
            '.md', '.py', '.txt', '.csv', '.html', '.css', '.js', '.json', '.xml',
            '.docx', '.jpg', '.jpeg', '.png', '.gif', '.webp', '.heic', '.heif'
        ]
        return file_extension in convertible_extensions
    
    def get_file_preview(self, file_path: str, max_chars=500) -> str:
        """Get file preview text"""
        try:
            if self._is_image_file(file_path):
                return f"[Image file: {os.path.basename(file_path)}]"
            elif self._is_audio_file(file_path):
                return f"[Audio file: {os.path.basename(file_path)}]"
            elif file_path.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read(max_chars)
                    if len(content) == max_chars:
                        content += "..."
                    return content
            else:
                return f"[Document: {os.path.basename(file_path)}]"
        except Exception as e:
            return f"[Error reading file: {e}]"

    def convert_to_pdf(self, file_path: str, output_path: str = None) -> str:
        """Convert various file types to PDF"""
        if output_path is None:
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = str(self.downloads_dir / f"{base_name}_converted_{timestamp}.pdf")

        file_extension = os.path.splitext(file_path)[1].lower()

        try:
            if file_extension == '.md':
                return self._convert_markdown_to_pdf(file_path, output_path)
            elif file_extension == '.py':
                return self._convert_code_to_pdf(file_path, output_path)
            elif file_extension == '.docx':
                return self._convert_docx_to_pdf(file_path, output_path)
            elif file_extension in ['.txt', '.csv', '.html', '.css', '.js', '.json', '.xml']:
                return self._convert_text_to_pdf(file_path, output_path)
            elif self._is_image_file(file_path):
                return self._convert_image_to_pdf(file_path, output_path)
            else:
                raise Exception(f"Unsupported file type for PDF conversion: {file_extension}")

        except Exception as e:
            raise Exception(f"Failed to convert {file_path} to PDF: {str(e)}")

    def _convert_markdown_to_pdf(self, file_path: str, output_path: str) -> str:
        """Convert Markdown file to PDF"""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Convert markdown to HTML
        html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'toc'])

        # Create PDF
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=20,
        )
        story.append(Paragraph(f"Converted from: {os.path.basename(file_path)}", title_style))
        story.append(Spacer(1, 12))

        # Convert HTML to PDF paragraphs (basic conversion)
        # Remove HTML tags and convert to plain text for simplicity
        import re
        text_content = re.sub('<[^<]+?>', '', html_content)
        paragraphs = text_content.split('\n\n')

        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), styles['Normal']))
                story.append(Spacer(1, 6))

        doc.build(story)
        return output_path

    def _convert_code_to_pdf(self, file_path: str, output_path: str) -> str:
        """Convert code file to PDF with syntax highlighting"""
        with open(file_path, 'r', encoding='utf-8') as f:
            code_content = f.read()

        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=20,
        )
        story.append(Paragraph(f"Code File: {os.path.basename(file_path)}", title_style))
        story.append(Spacer(1, 12))

        # Create code style
        code_style = ParagraphStyle(
            'Code',
            parent=styles['Normal'],
            fontName='Courier',
            fontSize=8,
            leftIndent=20,
            rightIndent=20,
            spaceAfter=6,
            borderWidth=1,
            borderColor='gray',
            borderPadding=5,
        )

        # Split code into lines and add to PDF
        lines = code_content.split('\n')
        for i, line in enumerate(lines, 1):
            # Escape special characters
            line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(f"{i:4d}: {line}", code_style))

        doc.build(story)
        return output_path

    def _convert_docx_to_pdf(self, file_path: str, output_path: str) -> str:
        """Convert DOCX file to PDF"""
        try:
            # Try using python-docx to extract text
            doc = docx.Document(file_path)

            pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            # Add title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=20,
            )
            story.append(Paragraph(f"Converted from: {os.path.basename(file_path)}", title_style))
            story.append(Spacer(1, 12))

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    story.append(Paragraph(paragraph.text, styles['Normal']))
                    story.append(Spacer(1, 6))

            pdf_doc.build(story)
            return output_path

        except Exception as e:
            # Fallback: try using LibreOffice if available
            return self._convert_with_libreoffice(file_path, output_path)

    def _convert_text_to_pdf(self, file_path: str, output_path: str) -> str:
        """Convert text file to PDF"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=20,
        )
        story.append(Paragraph(f"Text File: {os.path.basename(file_path)}", title_style))
        story.append(Spacer(1, 12))

        # Add content
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Replace line breaks with proper spacing
                para = para.replace('\n', '<br/>')
                story.append(Paragraph(para, styles['Normal']))
                story.append(Spacer(1, 6))

        doc.build(story)
        return output_path

    def _convert_image_to_pdf(self, file_path: str, output_path: str) -> str:
        """Convert image file to PDF"""
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        story = []

        # Add title
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=20,
        )
        story.append(Paragraph(f"Image: {os.path.basename(file_path)}", title_style))
        story.append(Spacer(1, 12))

        # Add image
        try:
            # Calculate image size to fit on page
            img = Image.open(file_path)
            img_width, img_height = img.size

            # Calculate scaling to fit on page (leave margins)
            max_width = 6 * inch  # 8.5" - 2.5" margins
            max_height = 8 * inch  # 11" - 3" margins

            scale = min(max_width / img_width, max_height / img_height, 1.0)

            scaled_width = img_width * scale
            scaled_height = img_height * scale

            # Add image to PDF
            pdf_image = ReportLabImage(file_path, width=scaled_width, height=scaled_height)
            story.append(pdf_image)

        except Exception as e:
            # If image can't be added, add error message
            story.append(Paragraph(f"Error loading image: {str(e)}", styles['Normal']))

        doc.build(story)
        return output_path

    def _convert_with_libreoffice(self, file_path: str, output_path: str) -> str:
        """Try to convert using LibreOffice (if available)"""
        try:
            # Try to use LibreOffice for conversion
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', os.path.dirname(output_path),
                file_path
            ]

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

            if result.returncode == 0:
                # LibreOffice creates PDF with same name as input file
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                generated_pdf = os.path.join(os.path.dirname(output_path), f"{base_name}.pdf")

                if os.path.exists(generated_pdf):
                    # Rename to desired output path
                    if generated_pdf != output_path:
                        shutil.move(generated_pdf, output_path)
                    return output_path

            raise Exception("LibreOffice conversion failed")

        except (subprocess.TimeoutExpired, FileNotFoundError, Exception):
            # LibreOffice not available or failed, use fallback
            return self._convert_text_to_pdf(file_path, output_path)

def convert_to_pdf(self, file_path: str, output_path: str = None) -> str:
    """Convert various file types to PDF"""
    if output_path is None:
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = str(self.downloads_dir / f"{base_name}_converted_{timestamp}.pdf")

    file_extension = os.path.splitext(file_path)[1].lower()

    try:
        if file_extension == '.md':
            return self._convert_markdown_to_pdf(file_path, output_path)
        elif file_extension == '.py':
            return self._convert_code_to_pdf(file_path, output_path)
        elif file_extension == '.docx':
            return self._convert_docx_to_pdf(file_path, output_path)
        elif file_extension in ['.txt', '.csv', '.html', '.css', '.js', '.json', '.xml']:
            return self._convert_text_to_pdf(file_path, output_path)
        elif self._is_image_file(file_path):
            return self._convert_image_to_pdf(file_path, output_path)
        else:
            raise Exception(f"Unsupported file type for PDF conversion: {file_extension}")

    except Exception as e:
        raise Exception(f"Failed to convert {file_path} to PDF: {str(e)}")

def _convert_markdown_to_pdf(self, file_path: str, output_path: str) -> str:
    """Convert Markdown file to PDF"""
    with open(file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Convert markdown to HTML
    html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'toc'])

    # Create PDF
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Add title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=20,
    )
    story.append(Paragraph(f"Converted from: {os.path.basename(file_path)}", title_style))
    story.append(Spacer(1, 12))

    # Convert HTML to PDF paragraphs (basic conversion)
    # Remove HTML tags and convert to plain text for simplicity
    import re
    text_content = re.sub('<[^<]+?>', '', html_content)
    paragraphs = text_content.split('\n\n')

    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para.strip(), styles['Normal']))
            story.append(Spacer(1, 6))

    doc.build(story)
    return output_path

def _convert_code_to_pdf(self, file_path: str, output_path: str) -> str:
    """Convert code file to PDF with syntax highlighting"""
    with open(file_path, 'r', encoding='utf-8') as f:
        code_content = f.read()

    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Add title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=20,
    )
    story.append(Paragraph(f"Code File: {os.path.basename(file_path)}", title_style))
    story.append(Spacer(1, 12))

    # Create code style
    code_style = ParagraphStyle(
        'Code',
        parent=styles['Normal'],
        fontName='Courier',
        fontSize=8,
        leftIndent=20,
        rightIndent=20,
        spaceAfter=6,
        borderWidth=1,
        borderColor='gray',
        borderPadding=5,
    )

    # Split code into lines and add to PDF
    lines = code_content.split('\n')
    for i, line in enumerate(lines, 1):
        # Escape special characters
        line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        story.append(Paragraph(f"{i:4d}: {line}", code_style))

    doc.build(story)
    return output_path

def _convert_docx_to_pdf(self, file_path: str, output_path: str) -> str:
    """Convert DOCX file to PDF"""
    try:
        # Try using python-docx to extract text
        doc = docx.Document(file_path)

        pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=20,
        )
        story.append(Paragraph(f"Converted from: {os.path.basename(file_path)}", title_style))
        story.append(Spacer(1, 12))

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                story.append(Paragraph(paragraph.text, styles['Normal']))
                story.append(Spacer(1, 6))

        pdf_doc.build(story)
        return output_path

    except Exception as e:
        # Fallback: try using LibreOffice if available
        return self._convert_with_libreoffice(file_path, output_path)

def _convert_text_to_pdf(self, file_path: str, output_path: str) -> str:
    """Convert text file to PDF"""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Add title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=20,
    )
    story.append(Paragraph(f"Text File: {os.path.basename(file_path)}", title_style))
    story.append(Spacer(1, 12))

    # Add content
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            # Replace line breaks with proper spacing
            para = para.replace('\n', '<br/>')
            story.append(Paragraph(para, styles['Normal']))
            story.append(Spacer(1, 6))

    doc.build(story)
    return output_path

def _convert_image_to_pdf(self, file_path: str, output_path: str) -> str:
    """Convert image file to PDF"""
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    story = []

    # Add title
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=20,
    )
    story.append(Paragraph(f"Image: {os.path.basename(file_path)}", title_style))
    story.append(Spacer(1, 12))

    # Add image
    try:
        # Calculate image size to fit on page
        img = Image.open(file_path)
        img_width, img_height = img.size

        # Calculate scaling to fit on page (leave margins)
        max_width = 6 * inch  # 8.5" - 2.5" margins
        max_height = 8 * inch  # 11" - 3" margins

        scale = min(max_width / img_width, max_height / img_height, 1.0)

        scaled_width = img_width * scale
        scaled_height = img_height * scale

        # Add image to PDF
        pdf_image = ReportLabImage(file_path, width=scaled_width, height=scaled_height)
        story.append(pdf_image)

    except Exception as e:
        # If image can't be added, add error message
        story.append(Paragraph(f"Error loading image: {str(e)}", styles['Normal']))

    doc.build(story)
    return output_path

def _convert_with_libreoffice(self, file_path: str, output_path: str) -> str:
    """Try to convert using LibreOffice (if available)"""
    try:
        # Try to use LibreOffice for conversion
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', os.path.dirname(output_path),
            file_path
        ]

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

        if result.returncode == 0:
            # LibreOffice creates PDF with same name as input file
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            generated_pdf = os.path.join(os.path.dirname(output_path), f"{base_name}.pdf")

            if os.path.exists(generated_pdf):
                # Rename to desired output path
                if generated_pdf != output_path:
                    shutil.move(generated_pdf, output_path)
                return output_path

        raise Exception("LibreOffice conversion failed")

    except (subprocess.TimeoutExpired, FileNotFoundError, Exception):
        # LibreOffice not available or failed, use fallback
        return self._convert_text_to_pdf(file_path, output_path)


class FileListWidget:
    """Widget for displaying and managing file lists"""

    def __init__(self, parent, file_manager: FileManager):
        self.parent = parent
        self.file_manager = file_manager
        self.files = []
        self.callbacks = {
            'on_file_select': None,
            'on_file_remove': None,
            'on_file_preview': None
        }

        self.create_widgets()
        self.setup_drag_drop()

    def create_widgets(self):
        """Create file list widgets"""
        # Main frame
        self.frame = tk.Frame(self.parent)

        # Toolbar
        toolbar = tk.Frame(self.frame)
        toolbar.pack(fill=tk.X, padx=5, pady=5)

        tk.Button(toolbar, text="Add Files",
                  command=self.add_files).pack(side=tk.LEFT, padx=(0, 5))
        tk.Button(toolbar, text="Clear All",
                  command=self.clear_files).pack(side=tk.LEFT)

        # File list
        list_frame = tk.Frame(self.frame)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Scrollable listbox
        from tkinter import scrolledtext
        self.file_listbox = tk.Listbox(list_frame, selectmode=tk.SINGLE, font=('Segoe UI', 9))
        scrollbar = tk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.file_listbox.yview)
        self.file_listbox.configure(yscrollcommand=scrollbar.set)

        self.file_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Bind events
        self.file_listbox.bind('<Double-Button-1>', self.on_double_click)
        self.file_listbox.bind('<<ListboxSelect>>', self.on_select)
        self.file_listbox.bind('<KeyPress-Delete>', self.on_delete_key)  # NEW: Delete key binding

        # Context menu
        self.context_menu = tk.Menu(self.file_listbox, tearoff=0)
        self.context_menu.add_command(label="Preview", command=self.preview_file)
        self.context_menu.add_separator()
        self.context_menu.add_command(label="Convert to PDF", command=self.convert_to_pdf)  # NEW
        self.context_menu.add_separator()
        self.context_menu.add_command(label="Remove", command=self.remove_file)
        self.context_menu.add_command(label="Show in Explorer", command=self.show_in_explorer)

        self.file_listbox.bind("<Button-3>", self.show_context_menu)

    def setup_drag_drop(self):
        """Setup drag and drop functionality"""
        try:
            # Make the listbox accept drops
            self.file_listbox.drop_target_register(tkdnd.DND_FILES)
            self.file_listbox.dnd_bind('<<Drop>>', self.on_drop)
        except:
            # tkinterdnd2 not available, skip drag and drop
            pass

    def on_drop(self, event):
        """Handle dropped files"""
        files = self.file_listbox.tk.splitlist(event.data)
        for file_path in files:
            if file_path not in [f['path'] for f in self.files]:
                valid, message = self.file_manager.validate_file(file_path)
                if valid:
                    file_info = self.file_manager.get_file_info(file_path)
                    self.files.append(file_info)
                    self.file_listbox.insert(tk.END, f"{file_info['name']} ({file_info['size_str']})")
                else:
                    messagebox.showerror("Invalid File", f"{os.path.basename(file_path)}: {message}")

        if self.callbacks['on_file_select']:
            self.callbacks['on_file_select'](self.get_selected_files())

    def add_files(self):
        """Add files to the list"""
        files = self.file_manager.select_files(multiple=True)
        for file_path in files:
            if file_path not in [f['path'] for f in self.files]:
                valid, message = self.file_manager.validate_file(file_path)
                if valid:
                    file_info = self.file_manager.get_file_info(file_path)
                    self.files.append(file_info)
                    self.file_listbox.insert(tk.END, f"{file_info['name']} ({file_info['size_str']})")
                else:
                    messagebox.showerror("Invalid File", f"{os.path.basename(file_path)}: {message}")

        if self.callbacks['on_file_select']:
            self.callbacks['on_file_select'](self.get_selected_files())

    def clear_files(self):
        """Clear all files from the list"""
        self.files.clear()
        self.file_listbox.delete(0, tk.END)

        if self.callbacks['on_file_select']:
            self.callbacks['on_file_select']([])

    def remove_file(self):
        """Remove selected file"""
        selection = self.file_listbox.curselection()
        if selection:
            index = selection[0]
            self.files.pop(index)
            self.file_listbox.delete(index)

            if self.callbacks['on_file_remove']:
                self.callbacks['on_file_remove'](index)

            if self.callbacks['on_file_select']:
                self.callbacks['on_file_select'](self.get_selected_files())

    def on_delete_key(self, event):
        """Handle Delete key press - NEW METHOD"""
        self.remove_file()

    def preview_file(self):
        """Preview selected file"""
        selection = self.file_listbox.curselection()
        if selection:
            file_info = self.files[selection[0]]
            if self.callbacks['on_file_preview']:
                self.callbacks['on_file_preview'](file_info)

    def show_in_explorer(self):
        """Show file in system explorer"""
        selection = self.file_listbox.curselection()
        if selection:
            file_path = self.files[selection[0]]['path']
            try:
                if os.name == 'nt':  # Windows
                    os.startfile(os.path.dirname(file_path))
                elif os.name == 'posix':  # macOS/Linux
                    os.system(f'open "{os.path.dirname(file_path)}"')
            except:
                messagebox.showerror("Error", "Could not open file location")

    def show_context_menu(self, event):
        """Show context menu"""
        if self.file_listbox.curselection():
            self.context_menu.post(event.x_root, event.y_root)

    def on_double_click(self, event):
        """Handle double-click on file"""
        self.preview_file()

    def on_select(self, event):
        """Handle file selection"""
        if self.callbacks['on_file_select']:
            self.callbacks['on_file_select'](self.get_selected_files())

    def get_selected_files(self) -> List[str]:
        """Get list of selected file paths"""
        return [f['path'] for f in self.files]

    def get_file_count(self) -> int:
        """Get number of files in list"""
        return len(self.files)

    def set_callback(self, event: str, callback):
        """Set callback for events"""
        if event in self.callbacks:
            self.callbacks[event] = callback

    def pack(self, **kwargs):
        """Pack the widget"""
        self.frame.pack(**kwargs)

    def grid(self, **kwargs):
        """Grid the widget"""
        self.frame.grid(**kwargs)

    def convert_to_pdf(self):
        """Convert selected file to PDF"""
        selection = self.file_listbox.curselection()
        if selection:
            file_info = self.files[selection[0]]
            file_path = file_info['path']

            if self.file_manager.can_convert_to_pdf(file_path):
                try:
                    # Show progress (optional)
                    import tkinter.messagebox as msgbox

                    # Convert to PDF
                    pdf_path = self.file_manager.convert_to_pdf(file_path)

                    # Ask if user wants to add PDF to file list
                    if msgbox.askyesno("PDF Created",
                                       f"PDF created successfully at:\n{pdf_path}\n\n"
                                       f"Would you like to add the PDF to your file list?"):
                        # Add PDF to file list
                        if pdf_path not in [f['path'] for f in self.files]:
                            pdf_info = self.file_manager.get_file_info(pdf_path)
                            self.files.append(pdf_info)
                            self.file_listbox.insert(tk.END, f"{pdf_info['name']} ({pdf_info['size_str']})")

                            if self.callbacks['on_file_select']:
                                self.callbacks['on_file_select'](self.get_selected_files())
                    else:
                        msgbox.showinfo("PDF Created", f"PDF saved to:\n{pdf_path}")

                except Exception as e:
                    msgbox.showerror("Conversion Error", f"Failed to convert file to PDF:\n{str(e)}")
            else:
                msgbox.showwarning("Not Supported", "This file type cannot be converted to PDF.")